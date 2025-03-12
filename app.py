import streamlit as st
from langchain_deepseek import ChatDeepSeek
from langchain.memory import ConversationBufferMemory
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import os
import docx
import fitz  # PyMuPDF for PDFs

from dotenv import load_dotenv
load_dotenv()



# Get API key from environment
api_key = os.getenv("DEEPSEEK_API_KEY")

if not api_key:
    st.error("API key not found! Make sure to set it in the .env file.")
else:
    os.environ["DEEPSEEK_API_KEY"] = api_key  # Optional: Explicitly set for LangChain


if 'key' not in st.session_state:
    st.session_state['key'] = 'value'




# Ensure data folder exists
DATA_FOLDER = "data"
os.makedirs(DATA_FOLDER, exist_ok=True)

st.title("AI-Powered Cover Letter Builder")





# --- File Upload Section ---
st.subheader("Upload your resume(s)")
st.write("Upload a document (TXT, PDF, DOCX)")
uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])

text = ""

if uploaded_file is not None:
    file_path = os.path.join(DATA_FOLDER, uploaded_file.name)

    # Save uploaded file permanently
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    st.success(f"File saved: `{file_path}`")

    # Extract text from the file
    file_type = uploaded_file.name.split(".")[-1]

    if file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif file_type == "pdf":
        pdf_document = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in pdf_document])
    elif file_type == "docx":
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

    st.text_area("Extracted Document Text", text, height=200)





# --- Load Existing Files ---
st.write("Load a saved file")
existing_files = [f for f in os.listdir(DATA_FOLDER) if f.endswith(("txt", "pdf", "docx"))]

selected_file = st.selectbox("Select a file to load:", ["None"] + existing_files)

if selected_file != "None":
    file_path = os.path.join(DATA_FOLDER, selected_file)
    file_type = selected_file.split(".")[-1]

    if file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif file_type == "pdf":
        pdf_document = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in pdf_document])
    elif file_type == "docx":
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

    st.text_area("Loaded Document Text", text, height=200)
    st.info(f"Loaded text from `{selected_file}`")


if 'cover_letter' not in st.session_state:
    st.session_state['cover_letter'] = ''


# Initialize LLM
llm = ChatDeepSeek(model="deepseek-chat", temperature=0.7)


print('App ready')



# Create a placeholder for the cover letter text area
cover_letter_placeholder = st.empty()

if not st.session_state.cover_letter:
    # --- Build the Cover Letter ---
    st.subheader("Build Your Cover Letter")
    job_description = st.text_area("Enter the job description:")

    default_prompt = ("Write a cover letter for the given position. Write up to three paragraphs, "
                      "using simple, personable, and heartfelt language.")
    user_input = st.text_input("Enter your prompt:", value=default_prompt)

    full_prompt = job_description + "\n" + user_input

    if st.button("Generate") and full_prompt.strip():
        prompt = f"Generate a professional cover letter based on the following details:\n\n{full_prompt}"
        response = llm.invoke([("human", prompt)])

        # Update session state cover letter
        st.session_state.cover_letter = response.content
        st.success("Cover Letter Generated!")

# Display the generated cover letter
if st.session_state.cover_letter:
    cover_letter_placeholder.text_area("Cover Letter", value=st.session_state.cover_letter, height=300, key="cover_letter_area")

    # --- Refine the Cover Letter ---
    st.subheader("Refine Your Cover Letter")
    edit_prompt = st.text_input("Give an instruction to improve the letter (e.g., 'Make it more concise'):")

    if st.button("Edit Cover Letter") and edit_prompt.strip():
        edit_template = PromptTemplate(
            input_variables=["existing_letter", "edit_instruction"],
            template=(
                "You are refining a cover letter. Here is the current version:\n\n{existing_letter}\n\n"
                "User's instruction: {edit_instruction}\n\nGenerate the improved version."
            )
        )
        # Combine the prompt template with the LLM using the runnable sequence operator
        edit_chain = edit_template | llm

        new_cover_letter = edit_chain.invoke({
            "existing_letter": st.session_state.cover_letter,
            "edit_instruction": edit_prompt
        })

        # Update session state with the new version
        st.session_state.cover_letter = new_cover_letter

        # Update the same placeholder with the updated cover letter
        cover_letter_placeholder.text_area("Cover Letter", value=st.session_state.cover_letter, height=300, key="cover_letter_area")

        st.success("Cover Letter Updated!")